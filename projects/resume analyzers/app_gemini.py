"""
Streamlit Resume, LinkedIn & GitHub Analyzer with Gemini API Integration
------------------------------------------------
Features:
- Upload resume (PDF, DOCX, TXT). Extract text and analyze.
- Paste job description / role keywords to compare against resume.
- Detect LinkedIn & GitHub URLs and give recommendations.
- Send resume + job description to Gemini (Generative API) for scoring, detailed feedback, and an updated resume in DOCX format following a recruiter-friendly format.
- Produce downloadable feedback report and updated resume (.docx).

Notes / Requirements:
- Uses: streamlit, requests, python-docx, PyPDF2, docx2txt
- Install (if needed): pip install streamlit requests python-docx PyPDF2 docx2txt
- You MUST provide a working Gemini/Generative API key and (optionally) endpoint details. This app uses a flexible HTTP call so you can adapt to your provider (Google Gemini / Vertex AI / other compatible endpoints).
- The app sends a structured prompt to the model asking for a JSON response with: overall_score (0-100), suggestions (list), updated_resume (full text in resume format), highlights (skills matched).

Security / Privacy:
- Resume text is sent to the external API you configure. Do not use sensitive personal data if you do not want it transmitted.

"""

import streamlit as st
import re
import io
import base64
from collections import Counter
import requests
import json
from docx import Document

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx2txt
except Exception:
    docx2txt = None

# --------------------------------------------------
# Configuration & Constants
# --------------------------------------------------

SKILLS_DB = [
    'python','pandas','numpy','scikit-learn','sklearn','tensorflow','pytorch',
    'sql','excel','power bi','tableau','matplotlib','seaborn','nlp','computer vision',
    'regression','classification','clustering','random forest','xgboost','lightgbm',
    'git','github','docker','aws','gcp','azure','bigquery','spark','hadoop'
]

SECTION_HEADERS = [
    'experience','education','projects','skills','certifications','summary','objective','publications','contact'
]

URL_PATTERN = re.compile(r'https?://\S+|www\.\S+')
EMAIL_PATTERN = re.compile(r'[\w\.-]+@[\w\.-]+')
PHONE_PATTERN = re.compile(r'(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{3}\)?[\s-]?|\d{3}[\s-]?)\d{3}[\s-]?\d{4}')
LINKEDIN_PATTERN = re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/\S+', re.IGNORECASE)
GITHUB_PATTERN = re.compile(r'(?:https?://)?(?:www\.)?github\.com/\S+', re.IGNORECASE)

# --------------------------------------------------
# Helper functions
# --------------------------------------------------

def extract_text_from_pdf(file_stream):
    if PyPDF2 is None:
        return ""  # PyPDF2 not installed
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = []
        for page in reader.pages:
            page_text = page.extract_text() or ''
            text.append(page_text)
        return '\n'.join(text)
    except Exception:
        return ""


def extract_text_from_docx(file_stream):
    if docx2txt is None:
        return ""
    try:
        tmp = file_stream.read()
        with open('/tmp/temp_resume.docx', 'wb') as f:
            f.write(tmp)
        text = docx2txt.process('/tmp/temp_resume.docx')
        return text
    except Exception:
        return ""


def extract_text_from_txt(file_stream):
    try:
        raw = file_stream.read()
        if isinstance(raw, bytes):
            raw = raw.decode('utf-8', errors='ignore')
        return raw
    except Exception:
        return ""


def clean_and_tokenize(text):
    t = text.lower()
    t = re.sub(r'[^a-z0-9\s]', ' ', t)
    tokens = [w for w in t.split() if len(w) > 1]
    return tokens


def detect_sections(text):
    text_lower = text.lower()
    found = {}
    for sec in SECTION_HEADERS:
        found[sec] = sec in text_lower
    return found


def generate_local_feedback(resume_text, jd_text=None):
    tokens = clean_and_tokenize(resume_text)
    sections = detect_sections(resume_text)
    skills_found = [s for s in SKILLS_DB if s in resume_text.lower()]

    urls = URL_PATTERN.findall(resume_text)
    emails = EMAIL_PATTERN.findall(resume_text)
    phones = PHONE_PATTERN.findall(resume_text)
    linkedin = LINKEDIN_PATTERN.findall(resume_text)
    github = GITHUB_PATTERN.findall(resume_text)

    length_words = len(tokens)
    length_score = 1.0 if 200 <= length_words <= 800 else max(0.2, min(1.0, length_words / 800))
    contact_score = 1.0 if (emails or phones) else 0.0
    section_score = sum(1 for v in sections.values() if v) / len(sections)
    skills_score = min(1.0, len(skills_found) / 8)

    overall = round((0.3 * length_score + 0.2 * contact_score + 0.2 * section_score + 0.3 * skills_score) * 100)

    feedback = {
        'overall_score': overall,
        'length_words': length_words,
        'length_score': round(length_score, 2),
        'contact_found': bool(emails or phones),
        'emails': emails,
        'phones': phones,
        'sections': sections,
        'section_score': round(section_score, 2),
        'skills_found': skills_found,
        'skills_score': round(skills_score, 2),
        'linkedin': linkedin,
        'github': github,
        'urls': urls,
    }

    if jd_text:
        resume_tokens = set(clean_and_tokenize(resume_text))
        jd_tokens = set(clean_and_tokenize(jd_text))
        common = resume_tokens.intersection(jd_tokens)
        if len(jd_tokens) == 0:
            jd_score = 0.0
        else:
            jd_score = round(min(len(common)/len(jd_tokens), 1.0), 3)
        feedback['job_match_score'] = int(jd_score*100)
        feedback['job_common_terms'] = len(common)
        feedback['job_skill_overlap'] = [s for s in SKILLS_DB if s in jd_tokens and s in resume_tokens]
    return feedback

# --------------------------------------------------
# Gemini API interaction
# --------------------------------------------------

def call_gemini_api(api_key, model_endpoint, prompt, timeout=60):
    """
    Flexible POST to a generative text endpoint. The default expectation is a JSON POST with {"prompt": "..."}
    If you use Google Generative Language API (v1beta2), set model_endpoint to:
    https://generativelanguage.googleapis.com/v1beta2/models/text-bison-001:generateText?key=YOUR_API_KEY

    If you use another provider, set endpoint accordingly and modify payload structure in this function.
    """
    headers = {
        'Content-Type': 'application/json'
    }
    payload = {
        'prompt': prompt,
        'max_output_tokens': 1024
    }

    try:
        # If model_endpoint already contains ?key=..., the api_key may be omitted
        if 'key=' in model_endpoint:
            resp = requests.post(model_endpoint, headers=headers, json=payload, timeout=timeout)
        else:
            # try adding api_key as query param
            url = model_endpoint
            if api_key:
                if '?' in url:
                    url = f"{url}&key={api_key}"
                else:
                    url = f"{url}?key={api_key}"
            resp = requests.post(url, headers=headers, json=payload, timeout=timeout)
        if resp.status_code != 200:
            return False, f"API Error {resp.status_code}: {resp.text}"
        return True, resp.json()
    except Exception as e:
        return False, str(e)

# Helper to build prompt for Gemini

def build_gemini_prompt(resume_text, jd_text=None, desired_format='recruiter_resume'):
    """
    Instruct the generative model to return a JSON object with these keys:
    - overall_score: integer 0-100
    - suggestions: list of bullet suggestions
    - updated_resume: full resume text in recruiter-friendly format (one-page for freshers)
    - highlights: list of matched skills

    The model is asked to output ONLY valid JSON. If the model outputs extra text, the app will try to extract the JSON block.
    """
    instruction = (
        "You are an expert career coach for Data Science & ML candidates.\n"
        "INPUTS:\n- resume_text: the candidate's resume plain text.\n"
        "- jd_text (optional): job description or role keywords.\n\n"
        "TASK:\n1) Evaluate the resume for recruiter-readiness and produce an overall_score (0-100).\n"
        "2) Provide a short list of suggestions to improve the resume (max 8 bullets).\n"
        "3) Produce an UPDATED_RESUME in a clear, recruiter-friendly format (one page for freshers, 1-2 pages for experienced).\n"
        "4) Provide a highlights list of matched skills found in the resume and JD.\n\n"
        "OUTPUT FORMAT:\nReturn ONLY a JSON object (no extra commentary) with keys: overall_score (int), suggestions (array of strings), updated_resume (string), highlights (array of strings).\n\n"
        "Make the updated_resume concise, use bullet points under each job/project with measurable impact, include contact header, skills section, projects, and education. If jd_text is provided, tailor the resume to include relevant keywords from the JD while preserving truthfulness.\n\n"
        "Begin JSON now.\n\n"
    )

    # Build body
    body = {
        'resume_text': resume_text,
        'jd_text': jd_text or ''
    }

    prompt = instruction + "INPUT_JSON:" + json.dumps(body)
    return prompt

# Utility: extract JSON from model output (robust)

def extract_json_from_text(text):
    # find first { ... }
    try:
        start = text.index('{')
        end = text.rindex('}')
        candidate = text[start:end+1]
        return json.loads(candidate)
    except Exception:
        # fallback: try to find line that starts with { and ends with }
        matches = re.findall(r'\{.*\}', text, flags=re.DOTALL)
        for m in matches:
            try:
                return json.loads(m)
            except Exception:
                continue
    return None

# --------------------------------------------------
# DOCX creation helper
# --------------------------------------------------

def create_docx_from_text(resume_text):
    doc = Document()
    lines = resume_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # If line looks like a heading (ALL CAPS or ends with ':'), make it bold
        if len(line) < 60 and (line.isupper() or line.endswith(':')):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --------------------------------------------------
# Streamlit UI
# --------------------------------------------------

st.set_page_config(page_title='Resume & Profile Analyzer (Gemini)', layout='wide')
st.title('Resume, LinkedIn & GitHub Analyzer (Gemini-powered)')
st.markdown('Upload your resume and optionally paste a job description. Provide your Gemini/Generative API key and endpoint, then click Analyze to get model-driven scoring, suggestions, and an updated resume (DOCX).')

col1, col2 = st.columns([1,2])
with col1:
    uploaded_file = st.file_uploader('Upload Resume', type=['pdf','docx','txt'])
    jd_text = st.text_area('Paste Job Description / Role Keywords (optional)', height=200)
    api_key = st.text_input('Generative API Key (Gemini/Vertex API Key)', type='password')
    model_endpoint = st.text_input('Model Endpoint URL', value='https://generativelanguage.googleapis.com/v1beta2/models/text-bison-001:generateText')
    analyze_btn = st.button('Analyze with Gemini')

with col2:
    st.info('Tips: If using Google Generative Language API, you can use the text-bison-001 model endpoint and provide the API key either in the query string or here. Adjust the Model Endpoint if you use a different provider.')
    st.subheader('Sample Skills (Data Science):')
    st.write(', '.join(SKILLS_DB))

resume_text = ''
if uploaded_file is not None:
    st.write('Uploaded file:', uploaded_file.name)
    if uploaded_file.name.lower().endswith('.pdf'):
        resume_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.lower().endswith('.docx'):
        resume_text = extract_text_from_docx(uploaded_file)
    else:
        resume_text = extract_text_from_txt(uploaded_file)

if analyze_btn:
    if not resume_text:
        st.error('No resume text found. Please upload a supported file with readable text.')
    elif not api_key:
        st.error('Please provide your Generative API Key to call Gemini/Generative API.')
    else:
        with st.spinner('Calling Gemini and preparing feedback...'):
            # Local quick analysis first
            local_feedback = generate_local_feedback(resume_text, jd_text)

            # Build prompt and call external API
            prompt = build_gemini_prompt(resume_text, jd_text)
            ok, resp = call_gemini_api(api_key, model_endpoint, prompt)
            if not ok:
                st.error(f'API call failed: {resp}')
            else:
                # Try to parse response JSON from model
                # Generative API responses differ by provider; attempt to extract text
                model_text = ''
                # If response has 'candidates' or 'output' fields adaptively pick
                if isinstance(resp, dict):
                    # Attempt to find text in common fields
                    if 'candidates' in resp and isinstance(resp['candidates'], list) and len(resp['candidates'])>0:
                        model_text = resp['candidates'][0].get('content', '') or json.dumps(resp['candidates'][0])
                    elif 'output' in resp:
                        model_text = resp['output'] if isinstance(resp['output'], str) else json.dumps(resp['output'])
                    else:
                        # fallback to stringifying full response
                        model_text = json.dumps(resp)
                else:
                    model_text = str(resp)

                # try extract JSON
                parsed = extract_json_from_text(model_text)
                if parsed is None:
                    st.warning('Could not extract structured JSON from model output. Showing raw output for review.')
                    st.code(model_text[:10000])
                else:
                    # Merge local feedback and model feedback where appropriate
                    overall = parsed.get('overall_score', local_feedback['overall_score'])
                    suggestions = parsed.get('suggestions', [])
                    updated_resume = parsed.get('updated_resume', '')
                    highlights = parsed.get('highlights', local_feedback['skills_found'])

                    st.metric('Overall Score (model)', f"{overall} / 100")
                    st.subheader('Top Suggestions (from model)')
                    for s in suggestions:
                        st.write('-', s)

                    st.subheader('Highlights (skills matched)')
                    st.write(', '.join(highlights) if highlights else 'None')

                    st.subheader('Updated Resume (Model-generated)')
                    st.text_area('Updated Resume', value=updated_resume, height=500)

                    # Create DOCX and provide download
                    docx_bytes = create_docx_from_text(updated_resume)
                    b64 = base64.b64encode(docx_bytes.getvalue()).decode()
                    href = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}"
                    st.markdown(f"[Download Updated Resume (DOCX)]({href})")

                    # Also produce a plain feedback report
                    report_lines = []
                    report_lines.append(f"Overall Score: {overall} / 100")
                    report_lines.append('\nSuggestions:')
                    report_lines.extend(suggestions)
                    report_text = '\n'.join(report_lines)
                    b = report_text.encode('utf-8')
                    href2 = f"data:file/txt;base64,{base64.b64encode(b).decode()}"
                    st.markdown(f"[Download Feedback Report]({href2})")

st.markdown('---')
st.caption('This app demonstrates how to combine local heuristics with a generative model (Gemini/Vertex/other). Adjust prompts and endpoint according to your provider. Do not send sensitive personal data unless you accept the privacy implications.')
